#!/usr/bin/env python3
"""End-to-end smoke test used by CI (and runnable locally).

Generates a tiny digital PDF with text + a table, converts it, and asserts the
resulting DOCX contains real, editable content (paragraphs and a table) — not a
flattened image. Exits non-zero on failure.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from pdf2docx_converter.config import ConvertSettings  # noqa: E402
from pdf2docx_converter.engine_local import convert_file  # noqa: E402


def _make_sample_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Smoke Test Document", fontsize=20, fontname="helv")
    page.insert_text((72, 110), "This paragraph must survive as editable text.",
                     fontsize=12, fontname="helv")
    shape = page.new_shape()
    for i in range(3):
        shape.draw_line((72, 150 + i * 24), (372, 150 + i * 24))
    for x in (72, 222, 372):
        shape.draw_line((x, 150), (x, 198))
    shape.finish(width=0.8)
    shape.commit()
    page.insert_text((80, 166), "Item", fontsize=11, fontname="helv")
    page.insert_text((230, 166), "Price", fontsize=11, fontname="helv")
    page.insert_text((80, 190), "Widget", fontsize=11, fontname="helv")
    page.insert_text((230, 190), "100", fontsize=11, fontname="helv")
    doc.save(path)
    doc.close()


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        pdf = tmp_path / "sample.pdf"
        docx = tmp_path / "sample.docx"

        _make_sample_pdf(pdf)
        convert_file(pdf, docx, ConvertSettings())

        assert docx.exists() and docx.stat().st_size > 0, "no DOCX produced"

        doc = Document(docx)
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert non_empty, "no editable paragraphs in output"
        assert doc.tables, "table was not preserved as an editable table"

        print(f"OK: {len(non_empty)} paragraph block(s), {len(doc.tables)} table(s).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
